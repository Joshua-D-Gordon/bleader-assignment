"""DOCX -> list[Block].

Walks the document body in true reading order (paragraphs and tables
interleaved), which python-docx's separate `.paragraphs` / `.tables` collections
do not preserve.

Heading detection is the load-bearing part. This file's headings do NOT use the
standard `Heading 1/2/3` styles — they use custom style IDs `heading10/20/30`
(display names "heading 10/20/30", but their real outline levels are 1/2/3) mixed
with real `Heading4`/`Heading5`. Parsing the level from the style *name* would
assign level 10 to a top-level heading. We instead read the style's `outlineLvl`
from the DOCX's styles.xml (heading_level = outlineLvl + 1), falling back to a
static styleId map for robustness. All heading styles in the target file are
covered by both paths.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..models import Block

# Fallback map: style_id (lowercased) -> heading level. Used only if outlineLvl
# cannot be resolved for a style.
_STYLE_ID_LEVEL = {
    "heading1": 1, "heading10": 1,
    "heading2": 2, "heading20": 2,
    "heading3": 3, "heading30": 3,
    "heading4": 4, "heading5": 5, "heading6": 6,
}


def _build_outline_level_map(doc: DocxDocument) -> dict[str, int]:
    """style_id -> heading level (1-based) via <w:outlineLvl> in styles.xml.

    outlineLvl is 0-based (0 = top level), so level = outlineLvl + 1. Only styles
    that actually declare an outline level are treated as headings.
    """
    levels: dict[str, int] = {}
    styles_el = doc.styles.element
    for style in styles_el.findall(qn("w:style")):
        style_id = style.get(qn("w:styleId"))
        if not style_id:
            continue
        # <w:pPr><w:outlineLvl w:val="N"/></w:pPr>
        outline = style.find(qn("w:pPr") + "/" + qn("w:outlineLvl"))
        if outline is not None:
            val = outline.get(qn("w:val"))
            if val is not None and val.isdigit():
                levels[style_id] = int(val) + 1
    return levels


def _iter_block_items(doc: DocxDocument):
    """Yield Paragraph and Table objects in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _heading_level(paragraph: Paragraph, outline_map: dict[str, int]) -> int | None:
    """Return the heading level for a paragraph, or None if it's body text."""
    style = paragraph.style
    if style is None:
        return None
    style_id = getattr(style, "style_id", None)
    if style_id and style_id in outline_map:
        return outline_map[style_id]
    # Fallbacks: static style-id map, then a loose name check.
    if style_id and style_id.lower() in _STYLE_ID_LEVEL:
        return _STYLE_ID_LEVEL[style_id.lower()]
    name = (style.name or "").strip().lower()
    if name.startswith("heading"):
        # "heading 2" -> 2, "heading 20" -> 2 (>=10 means custom 10/20/30 scheme)
        digits = re.findall(r"\d+", name)
        if digits:
            n = int(digits[0])
            return n // 10 if n >= 10 else n
    return None


def _table_to_markdown(table: Table) -> str:
    """Render a docx table as a GitHub-flavored markdown table."""
    rows = [[(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
            for row in table.rows]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]  # pad ragged rows
    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * width) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    return "\n".join([header, sep, body]) if body else "\n".join([header, sep])


def parse_docx(path: str | Path, doc_id: str) -> list[Block]:
    doc = Document(str(path))
    outline_map = _build_outline_level_map(doc)

    blocks: list[Block] = []
    section_path: list[str] = []
    order_index = 0
    para_n = 0
    table_n = 0

    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            para_n += 1
            text = item.text.strip()
            if not text:
                continue
            level = _heading_level(item, outline_map)
            if level is not None:
                # Trim the breadcrumb to the parent level, then append this heading.
                section_path = section_path[: level - 1] + [text]
                blocks.append(Block(
                    doc_id=doc_id, block_type="heading", heading_level=level,
                    text=text, section_path=section_path.copy(),
                    location=f"Paragraph {para_n}", order_index=order_index,
                ))
            else:
                blocks.append(Block(
                    doc_id=doc_id, block_type="paragraph", text=text,
                    section_path=section_path.copy(),
                    location=f"Paragraph {para_n}", order_index=order_index,
                ))
        elif isinstance(item, Table):
            table_n += 1
            md = _table_to_markdown(item)
            if not md:
                continue
            blocks.append(Block(
                doc_id=doc_id, block_type="table", text=md,
                section_path=section_path.copy(),
                location=f"Table {table_n}", order_index=order_index,
            ))
        order_index += 1

    return blocks


if __name__ == "__main__":  # standalone smoke test
    import sys
    from ..models import Block  # noqa
    p = sys.argv[1] if len(sys.argv) > 1 else "samples/FDS_PriceBook_Automation_V5.docx"
    bs = parse_docx(p, "B")
    headings = [b for b in bs if b.block_type == "heading"]
    tables = [b for b in bs if b.block_type == "table"]
    non_heading = [b for b in bs if b.block_type != "heading"]
    with_path = [b for b in non_heading if b.section_path]
    print(f"blocks={len(bs)} headings={len(headings)} tables={len(tables)}")
    print(f"non-heading blocks with section_path: "
          f"{len(with_path)}/{len(non_heading)} "
          f"({100*len(with_path)/max(len(non_heading),1):.1f}%)")
    for h in headings[:12]:
        print(f"  L{h.heading_level} {' > '.join(h.section_path)}")
